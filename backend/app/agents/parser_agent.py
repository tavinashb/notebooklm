import os
import re
import requests
from typing import List, Dict, Any, Optional
from pathlib import Path
import magic
import validators

# Document parsing imports
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

class DocumentChunk:
    def __init__(self, content: str, metadata: Dict[str, Any]):
        self.content = content
        self.metadata = metadata

class ParserAgent:
    """
    Agent responsible for document parsing and text extraction.
    Handles PDF, DOCX, TXT, and web URL inputs.
    """
    
    def __init__(self):
        self.max_chunk_size = 1000  # characters
        self.chunk_overlap = 200    # characters
        
    async def parse_document(self, file_path: str, file_type: str, url: Optional[str] = None) -> Dict[str, Any]:
        """
        Main entry point for document parsing.
        Returns parsed content with metadata and chunks.
        """
        try:
            if url and validators.url(url):
                return await self._parse_url(url)
            elif file_type.lower() == 'pdf':
                return await self._parse_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return await self._parse_docx(file_path)
            elif file_type.lower() == 'txt':
                return await self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "content": "",
                "chunks": [],
                "metadata": {}
            }
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF files and extract text content."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                page_metadata = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_content += f"\n\n--- Page {page_num + 1} ---\n\n{page_text}"
                    
                    page_metadata.append({
                        "page_number": page_num + 1,
                        "char_start": len(text_content) - len(page_text),
                        "char_end": len(text_content),
                        "text_length": len(page_text)
                    })
                
                # Create chunks
                chunks = self._create_semantic_chunks(text_content, {
                    "file_type": "pdf",
                    "total_pages": len(pdf_reader.pages),
                    "page_metadata": page_metadata
                })
                
                return {
                    "success": True,
                    "content": text_content.strip(),
                    "chunks": chunks,
                    "metadata": {
                        "file_type": "pdf",
                        "total_pages": len(pdf_reader.pages),
                        "total_chars": len(text_content),
                        "page_metadata": page_metadata
                    }
                }
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX files and extract text content."""
        try:
            doc = DocxDocument(file_path)
            text_content = ""
            paragraph_metadata = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    start_pos = len(text_content)
                    text_content += paragraph.text + "\n\n"
                    
                    # Check if paragraph is a heading
                    is_heading = paragraph.style.name.startswith('Heading') if paragraph.style else False
                    
                    paragraph_metadata.append({
                        "paragraph_index": i,
                        "char_start": start_pos,
                        "char_end": len(text_content),
                        "is_heading": is_heading,
                        "style": paragraph.style.name if paragraph.style else None
                    })
            
            # Create chunks
            chunks = self._create_semantic_chunks(text_content, {
                "file_type": "docx",
                "paragraph_metadata": paragraph_metadata
            })
            
            return {
                "success": True,
                "content": text_content.strip(),
                "chunks": chunks,
                "metadata": {
                    "file_type": "docx",
                    "total_paragraphs": len(doc.paragraphs),
                    "total_chars": len(text_content),
                    "paragraph_metadata": paragraph_metadata
                }
            }
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Parse TXT files and extract text content."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Create chunks
            chunks = self._create_semantic_chunks(text_content, {
                "file_type": "txt"
            })
            
            return {
                "success": True,
                "content": text_content,
                "chunks": chunks,
                "metadata": {
                    "file_type": "txt",
                    "total_chars": len(text_content),
                    "total_lines": len(text_content.split('\n'))
                }
            }
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
    
    async def _parse_url(self, url: str) -> Dict[str, Any]:
        """Parse web URLs and extract text content."""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "No title"
            
            # Extract main content
            text_content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks_text = '\n'.join(chunk for chunk in lines if chunk)
            
            # Create chunks
            chunks = self._create_semantic_chunks(chunks_text, {
                "file_type": "url",
                "url": url,
                "title": title_text
            })
            
            return {
                "success": True,
                "content": chunks_text,
                "chunks": chunks,
                "metadata": {
                    "file_type": "url",
                    "url": url,
                    "title": title_text,
                    "total_chars": len(chunks_text)
                }
            }
        except Exception as e:
            raise Exception(f"Error parsing URL: {str(e)}")
    
    def _create_semantic_chunks(self, text: str, base_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Create semantic chunks from text content.
        Prioritizes splitting by paragraphs and headings.
        """
        chunks = []
        
        # First, try to split by headings or major sections
        sections = self._split_by_headings(text)
        
        for section_idx, section in enumerate(sections):
            section_text = section["text"]
            section_metadata = section["metadata"]
            
            # If section is too large, split further
            if len(section_text) > self.max_chunk_size:
                sub_chunks = self._split_by_paragraphs(section_text)
                
                for sub_idx, sub_chunk in enumerate(sub_chunks):
                    chunk_metadata = {
                        **base_metadata,
                        **section_metadata,
                        "chunk_index": len(chunks),
                        "section_index": section_idx,
                        "sub_chunk_index": sub_idx,
                        "char_count": len(sub_chunk)
                    }
                    
                    chunks.append({
                        "content": sub_chunk.strip(),
                        "metadata": chunk_metadata
                    })
            else:
                chunk_metadata = {
                    **base_metadata,
                    **section_metadata,
                    "chunk_index": len(chunks),
                    "section_index": section_idx,
                    "char_count": len(section_text)
                }
                
                chunks.append({
                    "content": section_text.strip(),
                    "metadata": chunk_metadata
                })
        
        return chunks
    
    def _split_by_headings(self, text: str) -> List[Dict[str, Any]]:
        """Split text by headings and major sections."""
        # Look for heading patterns
        heading_patterns = [
            r'^#{1,6}\s+(.+)$',  # Markdown headings
            r'^([A-Z][A-Z\s]+)$',  # ALL CAPS headings
            r'^\d+\.\s+(.+)$',     # Numbered headings
        ]
        
        sections = []
        current_section = ""
        current_heading = None
        
        lines = text.split('\n')
        
        for line in lines:
            is_heading = False
            heading_text = None
            
            for pattern in heading_patterns:
                match = re.match(pattern, line.strip(), re.MULTILINE)
                if match:
                    is_heading = True
                    heading_text = match.group(1) if match.lastindex else line.strip()
                    break
            
            if is_heading and current_section.strip():
                # Save previous section
                sections.append({
                    "text": current_section.strip(),
                    "metadata": {
                        "section_header": current_heading,
                        "has_header": current_heading is not None
                    }
                })
                current_section = line + '\n'
                current_heading = heading_text
            else:
                current_section += line + '\n'
        
        # Add the last section
        if current_section.strip():
            sections.append({
                "text": current_section.strip(),
                "metadata": {
                    "section_header": current_heading,
                    "has_header": current_heading is not None
                }
            })
        
        return sections if sections else [{"text": text, "metadata": {"section_header": None, "has_header": False}}]
    
    def _split_by_paragraphs(self, text: str) -> List[str]:
        """Split large text into smaller chunks by paragraphs."""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) + 2 <= self.max_chunk_size:
                current_chunk += paragraph + '\n\n'
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + '\n\n'
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [text]